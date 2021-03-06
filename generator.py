'''
TODO: Contact selection for contact sheet
TODO: Read the contact data in from an easily edited word document so Jamie can add/remove contacts

TODO: (trivial) Centering accuracy into pdf dependent on physical size of characters. Python doens't distinguish between
for example, "i" and "B" for the purpose of spacing.

TODO: (Fancy) Auto-fill semicolon during time entry
'''

import os
import tkinter as tk
import openpyxl
import docx
import PyPDF2
import io
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


pdfmetrics.registerFont(TTFont('CalibriBd', 'calibrib.ttf'))
pdfmetrics.registerFont(TTFont('Calibri', 'calibri.ttf'))


SITE_NAME = 'The State of Stat Test Jail Prison'
SITE_CODE = 'TEST'
GROUPNO = 'PRN'+SITE_CODE
SITE_ADDRESS = '9999 Big Noodle Street, Testland TA 55555'
NAME_AND_CODE = SITE_NAME + ' - ' + SITE_CODE
REORDER_CUTOFF = '5:05PM'
NEWORDER_CUTOFF = '1:05PM'
TIMEZONE = 'ET'

# TODO: Ensure cutoff time input is valid
NEW_FILEPATH = 'testing\\'

PDF_XY_LANDSCAPE = (792, 612)
XY_PORTRAIT = (612, 792)

TEMPLATES = 'templates\\'

DOCUMENTS = []
CONTACTS = []

TIMEZONES = {
    'Eastern': '(ET)',
    'Central': '(CT)',
    'Mountain': '(MT)',
    'Pacific': '(PT)',
    'Alaska': '(AKT)',
    'Hawaii-Aleutian': '(HT)'
}


class GenApp(tk.Tk):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)

        self.title('New Site Document Generator')

        self.populate_contacts()
        self.populate_documents()
        self.create_widgets()

    def populate_contacts(self):
        CONTACTS.append(Contact('Chuck Allison', 'callison@diamondpharmacy.com', '2606'))
        CONTACTS.append(Contact('Dan Smith', 'dsmith@diamondpharmacy.com', '2627'))
        CONTACTS.append(Contact('Big Biff', 'dsmith@diamondpharmacy.com', '2627'))
        CONTACTS.append(Contact('Mortimer', 'dsmith@diamondpharmacy.com', '2627'))
        CONTACTS.append(Contact('Tweety', 'dsmith@diamondpharmacy.com', '2627'))

    def populate_documents(self):
        # Pdf
        DOCUMENTS.append(Document('Daily Drug 2016', DailyDrug))
        DOCUMENTS.append(Document('StatFormAscella249', StatFormAscella249))
        # Xlsx
        DOCUMENTS.append(Document('CONTROLLED STOCK NEW ORDER 71015', ControlledStockNewOrder))
        DOCUMENTS.append(Document('MEDICATION RETURN FORM', MedicationReturnForm))
        # Docx
        DOCUMENTS.append(Document('CONTROLLED STOCK REORDER', ControlledStockReorder))
        DOCUMENTS.append(Document('Fax Cover', FaxCover))
        DOCUMENTS.append(Document('NONCONTROLLED STOCK REORDER 10.2016', NonControlledStockReorder))
        DOCUMENTS.append(Document('Non-Formulary Request', NonFormularyRequest))
        DOCUMENTS.append(Document('PATIENT SPECIFIC REFILL FORM', PatientSpecificRefillForm))


    def create_widgets(self):

        self.name_label = tk.Label(self, text='Site name: ')
        self.name_entry = tk.Entry(self)

        self.code_label = tk.Label(self, text='Site code: ')
        self.code_entry = tk.Entry(self)

        self.address_label = tk.Label(self, text='Site address: ')
        self.address_entry = tk.Entry(self)

        self.cutoff_reorder_label = tk.Label(self, text='Re-order cutoff time: ')
        self.cutoff_reorder_entry = tk.Entry(self)
        self.cutoff_neworder_label = tk.Label(self, text='New-order cutoff time: ')
        self.cutoff_neworder_entry = tk.Entry(self)
        self.timezone_label = tk.Label(self, text='Time zone: ')
        self.tz_choice = tk.StringVar()
        self.tz_choice.set("Select TZ")
        self.timezone_menu = tk.OptionMenu(self, self.tz_choice, *TIMEZONES.items())

        self.contacts_label = tk.Label(self, text='Select contacts: ')
        self.contacts_listbox = tk.Listbox(self, selectmode='multiple')
        for c in CONTACTS:
            self.contacts_listbox.insert('end', c.name)

        self.documents_label = tk.Label(self, text='Select all documents you need: ')
        self.documents_listbox = tk.Listbox(self, selectmode='multiple', width=35)
        for d in DOCUMENTS:
            self.documents_listbox.insert('end', d.name)

        self.select_all_button = tk.Button(self, text='Select All', command=self.select_all_docs)
        self.deselect_all_button = tk.Button(self, text='Deselect All', command=self.deselect_all_docs)

        self.button = tk.Button(self, text='Give Birth', command=self.execute)

        current_row = 1
        self.name_label.grid(column=0, row=current_row, sticky='e')
        self.name_entry.grid(column=1, row=current_row, sticky='w')
        current_row += 1

        self.code_label.grid(column=0, row=current_row, sticky='e')
        self.code_entry.grid(column=1, row=current_row, sticky='w')
        current_row += 1

        self.address_label.grid(column=0, row=current_row, sticky='e')
        self.address_entry.grid(column=1, row=current_row, sticky='w')
        current_row += 1

        self.cutoff_reorder_label.grid(column=0, row=current_row, sticky='e')
        self.cutoff_reorder_entry.grid(column=1, row=current_row, sticky='w')
        current_row += 1

        self.cutoff_neworder_label.grid(column=0, row=current_row, sticky='e')
        self.cutoff_neworder_entry.grid(column=1, row=current_row, sticky='w')
        current_row += 1

        self.timezone_label.grid(column=0, row=current_row, sticky='e')
        self.timezone_menu.grid(column=1, row=current_row, sticky='w')
        current_row += 1

        self.documents_label.grid(column=0, row=current_row, sticky='e')
        self.documents_listbox.grid(column=1, row=current_row, rowspan=2)
        self.select_all_button.grid(column=0, row=current_row, sticky='s')
        current_row += 1
        self.deselect_all_button.grid(column=0, row=current_row, sticky='n')
        current_row += 1

        self.button.grid(column=0, row=current_row, columnspan=2)

    def get_input(self):
        global SITE_NAME, SITE_CODE, SITE_ADDRESS, REORDER_CUTOFF, NEWORDER_CUTOFF, NEW_FILEPATH, NAME_AND_CODE, GROUPNO
        SITE_NAME = self.name_entry.get()
        SITE_CODE = self.code_entry.get().upper()
        SITE_ADDRESS = self.address_entry.get()
        NAME_AND_CODE = SITE_NAME + ' - ' + SITE_CODE
        GROUPNO = 'PRN'+SITE_CODE
        print(self.tz_choice.get())
        TIMEZONE = ' ' + self.tz_choice.get()[-5:-3]
        print(TIMEZONE)
        REORDER_CUTOFF = self.cutoff_reorder_entry.get()+TIMEZONE
        NEWORDER_CUTOFF = self.cutoff_neworder_entry.get()+TIMEZONE

        NEW_FILEPATH = SITE_NAME + '\\'
        self.selected_docs = []
        for d in self.documents_listbox.curselection():
            self.selected_docs.append(self.documents_listbox.get(d))

    def select_all_docs(self):
        for i in range(self.documents_listbox.size()):
            self.documents_listbox.selection_set(i)

    def deselect_all_docs(self):
        for i in range(self.documents_listbox.size()):
            self.documents_listbox.selection_clear(i)

    def clear_fields(self):
        self.name_entry.delete(0, 'end')
        self.code_entry.delete(0, 'end')
        self.cutoff_reorder_entry.delete(0, 'end')
        self.cutoff_neworder_entry.delete(0, 'end')
        self.address_entry.delete(0, 'end')

    def make_dir(self):
        if not os.path.exists(SITE_NAME):
            os.makedirs(SITE_NAME)

    def execute(self):
        self.get_input()
        self.clear_fields()
        self.make_dir()
        for doc in DOCUMENTS:
            for d in self.selected_docs:
                if doc.name == d:
                    print(doc.name, 'matched')
                    doc.this_class().process()


class Contact:
    def __init__(self, name, email, extension):
        self.name = name
        self.email = email
        self.extension = extension


class Document:
    def __init__(self, name, this_class):
        self.name = name
        self.this_class = this_class


class Xlsx:
    def __init__(self):
        self.wb = openpyxl.load_workbook(TEMPLATES + self.filename)

    def process(self):
        self.replace()
        self.wb.save(NEW_FILEPATH + self.filename)

class ControlledStockNewOrder(Xlsx):
    def __init__(self):
        self.filename = 'CONTROLLED STOCK NEW ORDER 71015.xlsx'
        super().__init__()

    def replace(self):
        self.wb['Sheet1']['B4'].value = 'Facility Name/Code: ' + NAME_AND_CODE


class MedicationReturnForm(Xlsx):
    def __init__(self):
        self.filename = 'MEDICATION RETURN FORM.xlsx'
        super().__init__()

    def replace(self):
        self.wb['Sheet1']['A2'].value = 'FACILITY NAME: ' + SITE_NAME


class Docx:
    def show_contents(self):
        print(len(self.doc.paragraphs), ' paragraphs.')

        pNum = 1
        for p in self.doc.paragraphs:
            print('p', pNum, ': ', p.text)
            rNum = 1
            for r in p.runs:
                print ('run: ', rNum, r.text)
                rNum += 1
            pNum += 1

    def show_tables(self):
        print(len(self.doc.tables))

        tNum = 1
        for t in self.doc.tables:
            print('table: ', tNum, '\n')
            rNum=1
            for row in t.rows:
                print('row: ', rNum, '\n')
                cNum = 1
                for c in row.cells:
                    print('cell: ', cNum)
                    pNum = 1
                    for paragraph in c.paragraphs:
                        print('p: ', pNum)
                        runNum = 1
                        for run in paragraph.runs:
                            print('run: ', runNum, run.text)
                            runNum += 1
                        pNum += 1
                    cNum += 1
                rNum += 1
            tNum += 1

    def process(self):
        self.replace()
        self.doc.save(NEW_FILEPATH + self.filename)


class ControlledStockReorder(Docx):
    def __init__(self):
        self.filename = 'CONTROLLED STOCK REORDER.docx'
        self.doc = docx.Document(TEMPLATES + self.filename)

    def replace(self):
        for t in self.doc.tables:
            for row in t.rows:
                for c in row.cells:
                    for paragraph in c.paragraphs:
                        rNum = 0
                        for run in paragraph.runs:
                            if run.text == 'Name and Code':
                                run.text = NAME_AND_CODE
                            elif run.text == 'REFILL CUTOFF TIME ':
                                paragraph.runs[rNum+1].text = REORDER_CUTOFF
                                paragraph.runs[rNum+2].text = ''
                                paragraph.runs[rNum+3].text = ''
                            rNum += 1


class NonControlledStockReorder(Docx):
    def __init__(self):
        self.filename = 'NONCONTROLLED STOCK REORDER 10.2016.docx'
        self.doc = docx.Document(TEMPLATES + self.filename)

    def replace(self):
        for t in self.doc.tables:
            for row in t.rows:
                for c in row.cells:
                    for paragraph in c.paragraphs:
                        rNum = 0
                        for run in paragraph.runs:
                            if run.text == 'Name and Code':
                                run.text = NAME_AND_CODE
                            elif run.text == 'REFILL CUTOFF TIME ':
                                paragraph.runs[rNum + 1].text = REORDER_CUTOFF + '\n'
                            rNum += 1


class FaxCover(Docx):
    def __init__(self):
        self.filename = 'Fax Cover.docx'
        self.doc = docx.Document(TEMPLATES + self.filename)

    def replace(self):
        self.doc.paragraphs[2].runs[1].text = SITE_CODE[0]
        self.doc.paragraphs[2].runs[2].text = SITE_CODE[1]
        self.doc.paragraphs[2].runs[3].text = SITE_CODE[2:4]
        self.doc.paragraphs[13].runs[3].text = NAME_AND_CODE


class NonFormularyRequest(Docx):
    def __init__(self):
        self.filename = 'Non-Formulary Request.docx'
        self.doc = docx.Document(TEMPLATES + self.filename)

    def replace(self):
        self.doc.paragraphs[2].runs[3].text = NAME_AND_CODE


class PatientSpecificRefillForm(Docx):
    def __init__(self):
        self.filename = 'PATIENT SPECIFIC REFILL FORM.docx'
        self.doc = docx.Document(TEMPLATES + self.filename)

    def replace(self):
        for t in self.doc.tables:
            for row in t.rows:
                for c in row.cells:
                    for paragraph in c.paragraphs:
                        for run in paragraph.runs:
                            if run.text == 'Name and Code':
                                run.text = NAME_AND_CODE
                            elif run.text == 'REFILL CUTOFF TIME':
                                run.text = 'REFILL CUTOFF TIME ' + REORDER_CUTOFF


class Pdf:

    def merge(self):
        self.packet.seek(0)
        overlay_reader = PyPDF2.PdfFileReader(self.packet)
        self.page.mergePage(overlay_reader.getPage(0))

    def process(self):
        self.replace()
        self.merge()
        output_file = open(NEW_FILEPATH+self.output_filename, 'wb')
        self.output_writer.addPage(self.page)
        self.output_writer.write(output_file)
        output_file.close()
        self.base_file.close()

    def show_fields(self):
        fields = self.base_reader.getFields()
        print(fields)
        for k, v in fields.items():
            print(k, v)

class StatFormAscella249(Pdf):
    def __init__(self):
        self.base_file = open(TEMPLATES+'STATFormAscella-249.pdf', 'rb')
        self.base_reader = PyPDF2.PdfFileReader(self.base_file)
        self.page = self.base_reader.getPage(0)

        self.output_filename = 'STATFormAscella-249'+'-'+SITE_CODE+'.pdf'
        self.output_writer = PyPDF2.PdfFileWriter()

    def replace(self):

        self.packet = io.BytesIO()

        overlay_can = canvas.Canvas(self.packet, PDF_XY_LANDSCAPE)
        overlay_can.setFont('Helvetica', 11)
        overlay_can.drawString(88, 498, NAME_AND_CODE)
        overlay_can.drawString(84, 304, SITE_ADDRESS)
        overlay_can.drawString(644, 232, GROUPNO)
        overlay_can.save()

class DailyDrug(Pdf):
    def __init__(self):
        self.base_file = open(TEMPLATES+'Daily Drug 2016.pdf', 'rb')
        self.base_reader = PyPDF2.PdfFileReader(self.base_file)
        self.page = self.base_reader.getPage(0)

        self.output_filename = 'Daily Drug 2016.pdf'
        self.output_writer = PyPDF2.PdfFileWriter()

    def merge(self):
        self.packet.seek(0)
        overlay_reader = PyPDF2.PdfFileReader(self.packet)
        self.page.mergePage(overlay_reader.getPage(0))

    def replace(self):

        self.packet = io.BytesIO()

        overlay_can = canvas.Canvas(self.packet, PDF_XY_LANDSCAPE)
        overlay_can.setFont('Calibri', 11)
        overlay_can.drawString(648, 558, REORDER_CUTOFF)
        overlay_can.drawString(659, 531, NEWORDER_CUTOFF)
        overlay_can.drawString(512, 567, SITE_CODE)
        overlay_can.setFont('CalibriBd', 14)
        overlay_can.drawString(180, 565, self.center_name())
        overlay_can.save()

    def center_name(self):
        return f"{SITE_NAME:^54}"


DailyDrug().process()

app = GenApp()
app.mainloop()
